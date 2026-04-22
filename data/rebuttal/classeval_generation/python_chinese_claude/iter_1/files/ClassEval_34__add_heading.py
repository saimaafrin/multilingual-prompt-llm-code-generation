class _M:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.style import WD_STYLE_TYPE
    
    class WordDocument:
        def __init__(self):
            self.document = Document()
        
        def add_heading(self, heading, level=1):
            """
            向Word文档添加标题。
            :param heading: str，标题的文本。
            :param level: int，可选，标题的级别（1, 2, 3等；默认为1）。
            :return: bool，如果标题成功添加则返回True，否则返回False。
            """
            try:
                # 验证输入参数
                if not isinstance(heading, str):
                    return False
                
                if not isinstance(level, int) or level < 0:
                    return False
                
                # 添加标题到文档
                # level=0 表示 Title 样式
                # level=1-9 表示 Heading 1-9 样式
                self.document.add_heading(heading, level=level)
                
                return True
            except Exception as e:
                # 如果发生任何异常，返回False
                return False